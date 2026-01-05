import io
import re
import logging
import pandas as pd
from typing import List, Optional
from fastapi import UploadFile, HTTPException
from docx import Document

from app.core.config import settings

logger = logging.getLogger(__name__)

class SmartImportService:
    """
    Сервис для интеллектуального парсинга списков студентов из Excel, Word, TXT.
    """

    @staticmethod
    def normalize_name(raw_name: str) -> Optional[str]:
        """Очищает имя от мусора."""
        if not isinstance(raw_name, str):
            return None
        
        clean = re.sub(settings.NAME_SANITIZATION_REGEX, ' ', raw_name)
        clean = " ".join(clean.split())
        
        parts = [p.strip().capitalize() for p in clean.split() if len(p.strip()) > 1]
        
        if len(parts) >= 2:
            return " ".join(parts[:3])
        return None

    @staticmethod
    def _check_limit(count: int):
        if count > settings.MAX_STUDENTS_COUNT:
            raise HTTPException(
                status_code=400, 
                detail=f"Слишком много записей в файле (максимум {settings.MAX_STUDENTS_COUNT}). Разделите файл."
            )

    @classmethod
    async def parse_file(cls, file: UploadFile) -> List[dict]:
        # Читаем содержимое (размер уже проверен в endpoint)
        content = await file.read()
        filename = file.filename.lower()
        
        students = []

        try:
            if filename.endswith(('.xlsx', '.xls', '.csv')):
                students = cls._parse_excel(content, filename)
            elif filename.endswith('.docx'):
                students = cls._parse_docx(content)
            elif filename.endswith('.txt'):
                students = cls._parse_txt(content)
            else:
                raise HTTPException(status_code=400, detail="Неподдерживаемый формат файла")
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Import Error: {e}")
            raise HTTPException(status_code=400, detail=f"Ошибка чтения файла: {str(e)}")
            
        if not students:
             raise HTTPException(status_code=400, detail="Не удалось найти студентов в файле")

        return students

    @classmethod
    def _parse_excel(cls, content: bytes, filename: str) -> List[dict]:
        try:
            if filename.endswith('.csv'):
                df = pd.read_csv(io.BytesIO(content))
            else:
                df = pd.read_excel(io.BytesIO(content), header=None)
        except Exception as e:
            logger.warning(f"Failed to parse Excel with encoding detection: {e}")
            raise ValueError("Файл поврежден или имеет неверный формат")

        results = []
        
        # Стратегия 1: Одна колонка
        if df.shape[1] == 1:
            for raw in df.iloc[:, 0].dropna().astype(str):
                name = cls.normalize_name(raw)
                if name:
                    results.append({"full_name": name})
            
            cls._check_limit(len(results))
            return results
        
        # Стратегия 2: Поиск колонки
        first_row = df.iloc[0].astype(str).tolist()
        has_header = any(keyword in str(first_row).lower() for keyword in ['фио', 'фамилия', 'имя', 'студент', '№', 'no'])
        
        if has_header:
            df.columns = df.iloc[0]
            df = df.iloc[1:].reset_index(drop=True)
        
        fio_col = None
        for col in df.columns:
            sample = df[col].dropna().head(10).astype(str).tolist()
            fio_matches = sum(1 for x in sample if len(x.split()) >= 2 and sum(c.isalpha() for c in x) > len(x) * 0.7)
            if fio_matches >= 3:
                fio_col = col
                break
        
        if not fio_col:
            for col in reversed(df.columns.tolist()):
                if df[col].dtype == object:
                    fio_col = col
                    break
        
        if not fio_col:
            fio_col = df.columns[-1]
        
        for raw in df[fio_col].dropna().astype(str):
            name = cls.normalize_name(raw)
            if name:
                results.append({"full_name": name})
        
        cls._check_limit(len(results))
        return results

    @classmethod
    def _parse_docx(cls, content: bytes) -> List[dict]:
        doc = Document(io.BytesIO(content))
        names = []

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    for cell_text in reversed(cells):
                        name = cls.normalize_name(cell_text)
                        if name:
                            names.append(name)
                            break
            # Check limit inside loop to fail fast
            if len(names) > settings.MAX_STUDENTS_COUNT:
                 cls._check_limit(len(names))

        for para in doc.paragraphs:
            text = para.text.strip()
            if text and len(text) < 150:
                name = cls.normalize_name(text)
                if name:
                    names.append(name)
        
        seen = set()
        unique_names = []
        for n in names:
            if n not in seen:
                unique_names.append(n)
                seen.add(n)

        cls._check_limit(len(unique_names))
        return [{"full_name": n} for n in unique_names]

    @classmethod
    def _parse_txt(cls, content: bytes) -> List[dict]:
        for encoding in ['utf-8', 'cp1251', 'latin-1']:
            try:
                text = content.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        else:
            text = content.decode('utf-8', errors='ignore')
        
        names = []
        for line in text.split('\n'):
            line = line.strip()
            if line:
                name = cls.normalize_name(line)
                if name:
                    names.append(name)
        
        cls._check_limit(len(names))
        return [{"full_name": n} for n in names]